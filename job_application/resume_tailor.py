import os
from docx import Document
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv

load_dotenv()

def extract_text_from_docx(file_path):
    """Reads a docx file and extracts its text."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def tailor_resume_content(base_resume_text, job_description):
    """Uses LLM to rewrite the resume to fit the job description on 1 page."""
    # Using Langchain with OpenAI (fallback to Gemini if configured)
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7) # Ensure OPENAI_API_KEY is in .env

    prompt = PromptTemplate.from_template(
        """
        You are an expert career coach and professional resume writer.
        Your task is to rewrite the candidate's resume to perfectly align with the target job description.
        
        CRITICAL CONSTRAINT: The resulting resume content MUST be concise enough to fit strictly on 1 page.
        Prioritize the most relevant experiences, use high-impact action verbs, and bold key skills that match the job description.
        Brag about the candidate's fitness for the role, but keep the facts accurate to the original resume.
        
        Job Description:
        {job_description}
        
        Candidate's Base Resume:
        {base_resume}
        
        Output the rewritten resume in a clear, professional plain text format (Markdown is acceptable for structure like headers and bullets).
        """
    )
    
    chain = prompt | llm
    
    print("Generating tailored resume content...")
    response = chain.invoke({
        "job_description": job_description,
        "base_resume": base_resume_text
    })
    
    return response.content

def create_tailored_docx(content, output_path):
    """Writes the tailored content to a new docx file with basic formatting."""
    doc = Document()
    
    # Simple parsing to make headers and bullets
    for line in content.split('\n'):
        if line.startswith('##') or line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())
            
    # For a real implementation, you'd add logic here to adjust font sizes 
    # dynamically to ensure it fits strictly on 1 page.
    
    try:
        doc.save(output_path)
        print(f"Successfully saved tailored resume to {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

def process_resume(job):
    """Main pipeline for Phase 2."""
    base_resume_path = "resume.docx"
    
    # Create a dummy base resume if it doesn't exist for testing
    if not os.path.exists(base_resume_path):
        doc = Document()
        doc.add_heading('John Doe - Software Engineer', 0)
        doc.add_paragraph('Experienced software engineer with a background in Python, JavaScript, and building web applications.')
        doc.save(base_resume_path)
        print(f"Created a dummy base resume at {base_resume_path}")

    base_resume_text = extract_text_from_docx(base_resume_path)
    
    if not base_resume_text:
        return None
        
    tailored_text = tailor_resume_content(base_resume_text, job['description'])
    
    company_name = "".join([c for c in job['company'] if c.isalpha() or c.isdigit()]).rstrip()
    output_filename = f"resume_{company_name}.docx"
    
    create_tailored_docx(tailored_text, output_filename)
    return output_filename

if __name__ == "__main__":
    # Test execution
    sample_job = {
        "company": "TechCorp",
        "description": "Looking for a Python dev with LLM experience."
    }
    process_resume(sample_job)